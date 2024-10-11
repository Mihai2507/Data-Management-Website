from flask import Flask, render_template, request, redirect, url_for, jsonify, send_from_directory
from connect import DBConnection
from adaugare import adauga_inregistrari
from stergere import sterge_inregistrare
from actualizare import actualizeaza_inregistrare
from vizualizare import obtine_date
from markupsafe import Markup
from datetime import datetime
import pymysql
import os
import re
from docx import Document

app = Flask(__name__)

DOCUMENTS_FOLDER = 'documents/'
TEMPLATES_FOLDER = 'template_documents/'

os.makedirs(DOCUMENTS_FOLDER, exist_ok=True)
os.makedirs(TEMPLATES_FOLDER, exist_ok=True)

db_config = DBConnection()

def generate_paths(doc_name):
    doc_path = os.path.join(DOCUMENTS_FOLDER, doc_name)
    template_path = os.path.join(TEMPLATES_FOLDER, f'{os.path.splitext(doc_name)[0]}.txt')
    return doc_path, template_path

def load_replacements_from_template(template_path):
    replacements = {}
    if not os.path.exists(template_path):
        print(f"Template file does not exist: {template_path}")
        return replacements

    try:
        with open(template_path, 'r') as file:
            for line in file:
                if '=' in line:
                    old_string, new_string = line.strip().split('=', 1)
                    replacements[old_string.strip()] = new_string.strip()
    except Exception as e:
        print(f"Error reading template file: {e}")
    return replacements

def replace_strings_in_docx(doc_path, replacements):
    connection = None
    replacements_dict = {}

    try:
        connection = pymysql.connect(**db_config.get_connection_params())
        doc = Document(doc_path)
    except Exception as e:
        print(f"Error loading document or database connection: {e}")
        if connection:
            connection.close()
        return None

    for key, value in replacements.items():
        if '.' not in value:
            replacements_dict[key] = value
            continue

        table_name, field_name = value.split('.', 1)

        if table_name not in ['cursant', 'companie', 'comisie', 'curs']:
            print(f"Invalid table name: {table_name}")
            continue

        try:
            with connection.cursor() as cursor:
                cursor.execute(f"SHOW COLUMNS FROM {table_name}")
                columns = [row['Field'] for row in cursor.fetchall()]

                if field_name not in columns:
                    print(f"Column '{field_name}' not found in table '{table_name}'")
                    continue

                id_field = 'id'
                id_value = request.form.get(table_name)

                if not id_value:
                    print(f"No ID value provided for table '{table_name}'")
                    continue

                query = f"SELECT {field_name} FROM {table_name} WHERE {id_field} = %s"
                cursor.execute(query, (id_value,))
                result = cursor.fetchone()

                if result:
                    replacements_dict[key] = str(result[field_name])  # Convert to string
                else:
                    print(f"No result found for {table_name} ID {id_value}")

        except Exception as e:
            print(f"Error accessing database: {e}")
            if connection:
                connection.close()
            return None

    if connection:
        connection.close()

    print(f"Replacements dictionary: {replacements_dict}")

    for para in doc.paragraphs:
        for old_string, new_string in replacements_dict.items():
            if old_string in para.text:
                for run in para.runs:
                    if old_string in run.text:
                        run.text = run.text.replace(old_string, new_string)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old_string, new_string in replacements_dict.items():
                        if old_string in para.text:
                            for run in para.runs:
                                if old_string in run.text:
                                    run.text = run.text.replace(old_string, new_string)

    timestamp = datetime.now().strftime('%Y-%m-%d_%H-%M-%S')
    output_filename = f"{os.path.splitext(os.path.basename(doc_path))[0]}_{timestamp}.docx"
    output_path = os.path.join(DOCUMENTS_FOLDER, output_filename)

    try:
        doc.save(output_path)
    except Exception as e:
        print(f"Error: {e}")
        return None

    return output_path

def fetch_database_data():
    cursants = []
    companies = []
    comities = []
    courses = []

    connection = None

    try:
        connection = pymysql.connect(**db_config.get_connection_params())
        with connection.cursor() as cursor:
            cursor.execute("SELECT id, nume FROM cursant")
            cursants = cursor.fetchall()

            cursor.execute("SELECT id, nume FROM companie")
            companies = cursor.fetchall()

            cursor.execute("SELECT id, director FROM comisie")
            comities = cursor.fetchall()

            cursor.execute("SELECT id, nume_curs FROM curs")
            courses = cursor.fetchall()
    except Exception as e:
        print(f"Error: {e}")
    finally:
        if connection:
            connection.close()

    return cursants, companies, comities, courses


def extract_placeholders(document_content):
    placeholder_pattern = r'__\d+__'
    placeholders = re.findall(placeholder_pattern, document_content)
    return placeholders


def convert_to_html(doc):
    html = []
    for para in doc.paragraphs:
        style = para.style
        indent = style.paragraph_format.left_indent.pt if style.paragraph_format.left_indent else 0
        html.append(f'<p style="margin-left: {indent}px;">{para.text}</p>')

    for table in doc.tables:
        html.append('<table class="table">')
        for row in table.rows:
            html.append('<tr>')
            for cell in row.cells:
                html.append(f'<td>{cell.text}</td>')
            html.append('</tr>')
        html.append('</table>')

    return Markup(''.join(html))


@app.route('/', methods=['GET', 'POST'])
def home():
    return render_template('index.html')


@app.route('/select_document', methods=['GET', 'POST'])
def select_document():
    documents = [doc for doc in os.listdir(DOCUMENTS_FOLDER) if doc.endswith('.docx')]
    return render_template('generare_documente.html', documents=documents)



@app.route('/get_data', methods=['POST'])
def get_data():
    document_name = request.form['document_name']
    doc_path, template_path = generate_paths(document_name)

    try:
        doc = Document(doc_path)
    except Exception as e:
        return jsonify({'error': f"Error loading document: {e}"})

    document_content = "\n".join([para.text for para in doc.paragraphs])

    if os.path.exists(template_path):
        replacements = load_replacements_from_template(template_path)
    else:
        replacements = {}
        print(f"Template file not found: {template_path}")

    placeholders = extract_placeholders(document_content)
    manual_placeholders = [p for p in placeholders if p not in replacements]

    cursants, companies, comities, courses = fetch_database_data()

    return jsonify({
        'document_content': convert_to_html(doc),
        'manual_placeholders': manual_placeholders,
        'replacements': replacements,
        'cursants': cursants,
        'companies': companies,
        'comities': comities,
        'courses': courses
    })

@app.route('/fetch_options', methods=['POST'])
def fetch_options():
    table_name = request.form['table_name']

    if table_name == 'cursant':
        query = "SELECT id, nume AS name FROM cursant"
    elif table_name == 'companie':
        query = "SELECT id, nume AS name FROM companie"
    elif table_name == 'comisie':
        query = "SELECT id, director AS name FROM comisie"
    elif table_name == 'curs':
        query = "SELECT id, nume_curs AS name FROM curs"
    else:
        return jsonify({'error': 'Invalid table name'})

    connection_params = db_config.get_connection_params()
    connection = None

    try:
        connection = pymysql.connect(**connection_params)
        with connection.cursor() as cursor:
            cursor.execute(query)
            options = cursor.fetchall()
    except Exception as e:
        return jsonify({'error': f"Error fetching options: {e}"})
    finally:
        if connection:
            connection.close()

    return jsonify(options)

@app.route('/replace', methods=['POST'])
def replace():
    document_name = request.form['document_name']
    manual_replacements = {key: request.form[key] for key in request.form if key.startswith('__')}

    doc_path, template_path = generate_paths(document_name)

    if os.path.exists(template_path):
        template_replacements = load_replacements_from_template(template_path)
    else:
        template_replacements = {}

    replacements = template_replacements.copy()
    replacements.update(manual_replacements)

    output_path = replace_strings_in_docx(doc_path, replacements)

    if output_path:
        return send_from_directory(DOCUMENTS_FOLDER, os.path.basename(output_path))
    else:
        return jsonify({'error': 'Failed to generate the modified document.'})

@app.route('/preview', methods=['POST'])
def preview():
    document_name = request.form['document_name']
    doc_path, template_path = generate_paths(document_name)

    try:
        doc = Document(doc_path)
    except Exception as e:
        return jsonify({'error': f"Error loading document: {e}"})

    document_content = convert_to_html(doc)

    # Extract placeholders to determine if they are manual or not
    placeholders = extract_placeholders(document_content)
    manual_placeholders = [p for p in placeholders if p not in load_replacements_from_template(template_path)]

    return jsonify({
        'document_content': document_content,
        'manual_placeholders': manual_placeholders
    })

@app.route('/generate', methods=['POST'])
def generate():
    document_name = request.form['document_name']
    manual_replacements = {key: request.form[key] for key in request.form if key.startswith('__')}

    doc_path, template_path = generate_paths(document_name)

    if os.path.exists(template_path):
        template_replacements = load_replacements_from_template(template_path)
    else:
        template_replacements = {}

    replacements = template_replacements.copy()
    replacements.update(manual_replacements)

    output_path = replace_strings_in_docx(doc_path, replacements)

    if output_path:
        return jsonify({
            'file_url': f'/download/{os.path.basename(output_path)}',
            'preview_content': convert_to_html(Document(output_path))
        })
    else:
        return jsonify({'error'})


@app.route('/download/<filename>')
def download_file(filename):
    return send_from_directory(DOCUMENTS_FOLDER, filename)


@app.route('/adaugare', methods=['GET', 'POST'])
def adaugare():
    if request.method == 'POST':
        form_type = request.form['form_type']
        if form_type == 'cursant':
            cursant_data = {
                'nume': request.form['cursantNume'],
                'cnp': request.form['cursantCNP'],
                'judet_sector': request.form['cursantJudetSector'],
                'localitate': request.form['cursantLocalitate'],
                'prenume_mama': request.form['cursantParinteMama'],
                'prenume_tata': request.form['cursantParinteTata'],
                'strada':request.form['cursantStrada'],
                'numar_strada': request.form['cursantNumarStrada'],
                'bloc': request.form['cursantBloc'],
                'scara_bloc': request.form['cursantScaraBloc'],
                'nr_apartament': request.form['cursantNrApartament'],
                'nr_telefon': request.form['cursantNrTelefon'],
                'anul_nasterii': request.form['cursantAnulNasterii'],
                'luna_nasterii': request.form['cursantLunaNasterii'],
                'ziua_nasterii': request.form['cursantZiuaNasterii'],
                'locul_nasterii': request.form['cursantLocNastere']
            }
            adauga_inregistrari(cursant_data, 'cursant')

        elif form_type == 'companie':
            companie_data = {
                'nume': request.form['companieNume'],
                'localitate': request.form['companieLocalitate'],
                'judet': request.form['companieJudet']
            }
            adauga_inregistrari(companie_data, 'companie')

        elif form_type == 'comisie':
            comisie_data = {
                'director': request.form['comisieDirector'],
                'secretar': request.form['comisieSecretar'],
                'presedinte': request.form['comisiePresedinte']
            }
            adauga_inregistrari(comisie_data, 'comisie')

        elif form_type == 'curs':
            curs_data = {
                'nume_curs': request.form['cursNume'],
                'descriere': request.form['cursDescriere'],
                'cor': request.form['cursCor'],
                'nr_registru': request.form['cursNrRegistru'],
                'data_incepere': request.form['cursDataIncepere'],
                'data_finalizare': request.form['cursDataFinalizare'],
                'durata': request.form['cursDurata'],
                'anul_examinarii': request.form['cursAnulExaminarii'],
                'luna_examinarii': request.form['cursLunaExaminarii'],
                'ziua_examinarii': request.form['cursZiuaExaminarii']
            }
            adauga_inregistrari(curs_data, 'curs')

        return redirect(url_for('adaugare'))
    return render_template('adaugare_date.html')


@app.route('/actualizare_date', methods=['GET', 'POST'])
def actualizare_date():
    if request.method == 'POST':
        table = request.form.get('table')
        column = request.form.get('column')
        old_value = request.form['old_value']
        new_value = request.form.get('new_value')
        if not table or not column or not new_value:
            return "Error: Missing data in form", 400

        actualizeaza_inregistrare(table, column, old_value, new_value)
        return redirect(url_for('actualizare_date'))
    return render_template('actualizare_date.html')


@app.route('/stergere_date', methods=['GET', 'POST'])
def stergere_date():
    if request.method == 'POST':
        table = request.form['table']
        column = request.form['column']
        value = request.form['value']
        sterge_inregistrare(table, column, value)
        return redirect(url_for('stergere_date'))
    return render_template('stergere_date.html')


@app.route('/vizualizare_date')
def vizualizare_date():
    return render_template('vizualizare_date.html')


@app.route('/api/cursanti')
def api_cursanti():
    date_cursanti = obtine_date('cursant')
    return jsonify(date_cursanti)


@app.route('/api/companii')
def api_companii():
    date_companii = obtine_date('companie')
    return jsonify(date_companii)


@app.route('/api/comisii')
def api_comisii():
    date_comisii = obtine_date('comisie')
    return jsonify(date_comisii)


@app.route('/api/cursuri')
def api_cursuri():
    date_cursuri = obtine_date('curs')
    return jsonify(date_cursuri)


if __name__ == '__main__':
    app.run(debug=True)
